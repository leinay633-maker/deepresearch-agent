from __future__ import annotations

import asyncio
import json
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

from deepresearch_agent.config import Settings
from deepresearch_agent.ingest_corpus import ingest_directory
from deepresearch_agent.rag import LocalRagRetriever


def test_ingest_directory_writes_local_corpus_jsonl(tmp_path: Path) -> None:
    docs = tmp_path / "docs"
    docs.mkdir()
    (docs / "agent.md").write_text(
        "---\ntags: [agent]\n---\n# Agent Notes\n\nPlanner evidence and citation grounding.",
        encoding="utf-8",
    )
    (docs / "run-control.txt").write_text(
        "Worker leases keep queued runs from double execution.",
        encoding="utf-8",
    )
    _write_docx(docs / "report.docx")
    _write_pdf(docs / "brief.pdf")
    (docs / ".obsidian").mkdir()
    (docs / ".obsidian" / "workspace.json").write_text("do not ingest", encoding="utf-8")
    (docs / "empty.md").write_text("   \n", encoding="utf-8")
    output = tmp_path / "corpus.jsonl"

    summary = ingest_directory(docs, output)
    rows = [json.loads(line) for line in output.read_text(encoding="utf-8").splitlines()]

    assert summary.document_count == 4
    assert summary.skipped_count == 1
    assert [row["id"] for row in rows] == [
        "local-agent",
        "local-brief",
        "local-report",
        "local-run-control",
    ]
    assert rows[0]["title"] == "Agent Notes"
    assert "tags:" not in rows[0]["content"]
    assert rows[0]["metadata"]["source_path"] == "agent.md"
    assert "PDF evidence extraction works" in rows[1]["content"]
    assert rows[1]["metadata"]["ingest_format"] == "pdf"
    assert "DOCX private corpus citation text" in rows[2]["content"]
    assert rows[2]["metadata"]["ingest_format"] == "docx"
    assert rows[3]["title"] == "run control"
    assert rows[3]["url"].startswith("file:///")


def test_ingested_corpus_can_feed_local_retriever(tmp_path: Path) -> None:
    docs = tmp_path / "docs"
    docs.mkdir()
    (docs / "worker.md").write_text(
        "# Worker Queue\n\nQueued runs use request_json and worker lease claims.",
        encoding="utf-8",
    )
    output = tmp_path / "corpus.jsonl"
    ingest_directory(docs, output)
    retriever = LocalRagRetriever(
        corpus_path=output,
        settings=Settings(local_retrieval_mode="keyword"),
    )

    results = asyncio.run(retriever.retrieve("worker lease request_json", max_results=1))

    assert results[0].title == "Worker Queue"
    assert results[0].metadata["local_doc_id"] == "local-worker"


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Docx Report", level=1)
    document.add_paragraph("DOCX private corpus citation text.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "field"
    table.cell(0, 1).text = "value"
    document.save(path)


def _write_pdf(path: Path) -> None:
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 720, "PDF evidence extraction works.")
    pdf.save()
